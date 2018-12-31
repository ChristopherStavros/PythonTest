'''
pip install python-docx # import docx
'''
import docx
import os, sys
#########################################################
#Set script path and remove trailing \ if path is c:\
#########################################################
scriptPath = os.path.dirname(sys.argv[0])
if scriptPath[-1:]=='\\':
    scriptPath = scriptPath.strip('\\')
#########################################################
'''
import requests
#get file and save it in variable --> res
res = requests.get('http://autbor.com/demo.docx')
res.raise_for_status()

#DOWNLOAD FILE!
playfile = open('%s\\demo.docx' % (scriptPath), 'wb') # open file in wb --- write-binary mode
for chunk in res.iter_content(100000):  # specify number of bytes per chunk
    playfile.write(chunk)

playfile.close()
'''
os.chdir(scriptPath)
print(os.getcwd())
#########################################################
'''
docx structure
---'document' object contains 'paragraph' objects
---'paragraph' object contains 'run' objects
'''
d = docx.Document('demo.docx')
#print(type(d)) # <class 'docx.document.Document'>
#print(type(d.paragraphs)) # <class 'list'>
#print(d.paragraphs) # [<docx.text.paragraph.Paragraph object at 0x000002A345C7CC50>, <docx.text.paragraph.Paragraph object at 0x000002A345C7CF28>, <docx.text.paragraph.Paragraph object at 0x000002A345C7CF98>, <docx.text.paragraph.Paragraph object at 0x000002A345C7CE80>, <docx.text.paragraph.Paragraph object at 0x000002A345C93128>, <docx.text.paragraph.Paragraph object at 0x000002A345C93198>, <docx.text.paragraph.Paragraph object at 0x000002A345C930F0>]

#print text
print(d.paragraphs[0].text)
for i in d.paragraphs:
    #print(type(i)) # <class 'docx.text.paragraph.Paragraph'>
    print(i.text)





